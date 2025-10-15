
from docx import Document
import os
import uuid
from typing import List, Dict
from xml.etree import ElementTree as ET

NS = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
}

def is_list_paragraph(paragraph):
    # проверяем наличие numPr в XML параграфа
    p = paragraph._p
    numPr = p.find('.//w:numPr', NS)
    return numPr is not None

def extract_images(doc: Document, out_dir: str):
    os.makedirs(out_dir, exist_ok=True)
    saved = []
    for rel in doc.part.rels.values():
        try:
            tp = rel.target_part
            ctype = getattr(tp, "content_type", "")
            if "image" in ctype:
                img = tp.blob
                ext = ctype.split('/')[-1]
                name = f"img_{uuid.uuid4().hex}.{ext}"
                path = os.path.join(out_dir, name)
                with open(path, "wb") as f:
                    f.write(img)
                saved.append(path)
        except Exception:
            # игнорировать нерелевантные связи
            continue
    return saved

def read_docx_sections(path: str) -> List[Dict]:
    doc = Document(path)
    base_name = os.path.basename(path)
    sections = []
    # initialize current section with metadata
    current = {
        "title": "ROOT",
        "content": "",
        "type": "section",
        "source_file": base_name,
        "section_id": uuid.uuid4().hex
    }

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue

        # Heading styles: Heading 1, Heading 2, etc.
        if style_name and style_name.lower().startswith("heading"):
            # store previous if non-empty
            if current and (current.get("content") or current.get("title")):
                sections.append(current)
            # start new section with metadata
            current = {
                "title": text,
                "content": "",
                "type": "heading",
                "style": style_name,
                "source_file": base_name,
                "section_id": uuid.uuid4().hex
            }
            continue

        # lists
        if is_list_paragraph(para):
            current["content"] += "\n- " + text
            continue

        # normal paragraph
        current["content"] += ("\n" if current["content"] else "") + text

    # append last
    if current and (current["content"] or current["title"]):
        sections.append(current)

    # tables
    for table in doc.tables:
        rows = []
        for r in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in r.cells]
            rows.append("| " + " | ".join(cells) + " |")
        sections.append({
            "title": "TABLE",
            "content": "\n".join(rows),
            "type": "table",
            "source_file": base_name,
            "section_id": uuid.uuid4().hex
        })

    # images (saved to /data/images)
    imgs = extract_images(doc, out_dir=os.path.join(os.path.dirname(path), "images"))
    for p in imgs:
        sections.append({
            "title": "IMAGE",
            "content": p,
            "type": "image",
            "source_file": base_name,
            "section_id": uuid.uuid4().hex
        })

    return sections

if __name__ == "__main__":
    import sys
    path = sys.argv[1]
    secs = read_docx_sections(path)
    for s in secs[:10]:
        print("=== ", s["title"])
        print(s["content"][:300])
        print()
